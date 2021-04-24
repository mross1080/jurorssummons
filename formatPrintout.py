from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import subprocess
import json
import datetime
import random


def testPrint():
    print("Connected to formatter")


answer_lookup = {
    "en": {
        "a1": {
            "1": "do",
            "2": "have",
            "Yes!": "do",
            "No!": "have"
        },
        "a3": {
            "1": "fully Implicated",
            "2": "modestly implicated",
            "3": "slightly implicated",
            "4": "not implicated at all"
        },
        "a2": {
            "1": "how you see and identify yourself",
            "2": "how others see and identify you",
            "3": "how you see yourself and how others see you",
            "4": "nothing"
        },
        "sugarIntake": {
            "1": "high",
            "2": "medium",
            "3": "low",
            "4": "none"
        }
    },
    "es": {
        "a1": {
            "1": "conoce o ha oído",
            "2": "no conoce y no ha oído",
            "Yes!": " conoce o ha oído",
            "No!": " no conoce y no ha oído"
        },
        "a2": {
            "1": "totalmente implicado",
            "2": "modestamente implicado",
            "3": "ligeramente implicado",
            "4": "para nada implicado"
        },
        "a3": {
            "1": "cómo te ves e identificas",
            "2": "como otres te ven e identifican",
            "3": "todo lo anterior",
            "4": "ninguna de las anteriores"
        },
        "sugarIntake": {
            "1": "alto",
            "2": "medio",
            "3": "bajo",
            "4": "ninguno"
        }
    }
}


zipcode_data_lookup = {}


with open('/home/pi/zipcode_data.txt') as json_file:
    zipcode_data_lookup = json.load(json_file)


def formatDocument(userInfo):
    print("Starting Custom Print Job")

    doc = Document("/home/pi/CivilReviewENTemplate.docx")
    lang = userInfo["lang"]
    if userInfo["lang"] == "es":
        doc = Document("/home/pi/CivilReviewESTemplate.docx")

    zipcode = userInfo["zipcode"]
    if zipcode not in zipcode_data_lookup.keys():
        zipcode = "11222"
    data_for_zip = zipcode_data_lookup[zipcode]
    print(data_for_zip)
    styles = doc.styles
    style = styles.add_style('Insertion', WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles['Insertion']
    font = style.font
    font.name = 'Helvetica'
    font.size = Pt(18)

    print("Creating Document using this info")
    print(userInfo)
    random_data_points = random.sample(range(1, len(data_for_zip["en"])), 3)
    for paragraph in doc.paragraphs:
        # print(paragraph.text)

        if '[DATE]' in paragraph.text:
            paragraph.style = 'Insertion'
            paragraph.text =f"{datetime.datetime.now():%m-%d-%Y}" 

        if '[NAME]' in paragraph.text:
            paragraph.style = 'Insertion'
            paragraph.text = 'Applicant: {}'.format(userInfo["userName"])
            if lang == "es":
                paragraph.text = 'Solicitante: {}'.format(userInfo["userName"])

        if '[QUALIFYSTATUS]' in paragraph.text:
            paragraph.style = 'Insertion'
            paragraph.text = 'Result : Disqualify'
            if lang == "es":
                paragraph.text = "Resultado : No Califica"

        if '[Q1Part1]' in paragraph.text or "[Q1]" in paragraph.text:   
            q1Answer = answer_lookup[lang]["a1"][userInfo["a1"]]
            q1Part1Answer = "[do]" if q1Answer == "do" else "[don't]"
            q1Part2Answer = "[not]" if q1Answer != "do" else ""


            if lang == "en":
                paragraph.text = "You {} know or have {} heard of any of the suspects and witnesses.  To be an impartial reviewer you can not know or have heard of the suspects and witnesses.".format(
                    q1Part1Answer, q1Part2Answer)
            else:
                q1Answer = answer_lookup[lang]["a1"][userInfo["a1"]]
                paragraph.text = "Usted [{}] de los testigos o los sospechosos.  Para ser un evaluador imparcial, no puede ni conocer, ni haber oído hablar de los sospechosos y testigos.".format(q1Answer)

        if '[Q2]' in paragraph.text:
            questionTwoAnswer = answer_lookup[lang]["a3"][userInfo["a3"]]

            paragraph.text = "You think the U.S. is [{}] in the racial constructs of the D.R..  To be an impartial reviewer you would have to recognize that the U.S.A. is a cultural, economic, and political hegemonic imperial power.".format(
                questionTwoAnswer)
            if lang == "es":
                paragraph.text = "Usted cree que los Estados Unidos está [{}] en las construcciones raciales de R.D..  Para ser un evaluador imparcial, tendría que reconocer que los Estados Unidos es una potencia imperial hegemónica cultural, económica y política.".format(questionTwoAnswer)

        if '[Q3]' in paragraph.text:
            questionThreeAnswer = answer_lookup[lang]["a2"][userInfo["a2"]]
            paragraph.text = "For you, [{}] affects your material conditions.  To be an impartial reviewer you would have to see 'othering' mechanisms as tools of control.  ".format(questionThreeAnswer)
            if lang == "es":
                paragraph.text = "Para usted, [{}] afecta sus condiciones materiales.  Para ser un evaluador imparcial, tendría que ver los mecanismos 'de alteridad' como herramientas de control.".format(questionThreeAnswer)

        if '[Q4]' in paragraph.text:
            questionFourAnswer = answer_lookup[lang]["sugarIntake"][userInfo["sugarIntake"]]
            paragraph.text = "Your weekly sugar intake is {}.  To be an impartial reviewer would not consume any sugar.".format(
                answer_lookup[lang]["sugarIntake"][userInfo["sugarIntake"]])
            if lang == "es":
                paragraph.text = "Su consumo semanal de azúcar es [{}].  Para ser un evaluador imparcial, no debería consumir azúcar.".format(questionFourAnswer)

        if '[XX%]' in paragraph.text:
            print("Setting Random Value 1")
            paragraph.style = 'Insertion'

            # print(data_for_zip[random_data_points[0]])
            paragraph.text = data_for_zip[lang][random_data_points[0]]


        if '[YY%]' in paragraph.text or '[$YY]' in paragraph.text:
            print("Setting Random Value 2")
            paragraph.style = 'Insertion'
            paragraph.text = data_for_zip[lang][random_data_points[1]]

        if '[ZZ%]' in paragraph.text:
            print("Setting Random Value 3")
            paragraph.style = 'Insertion'
            paragraph.text = data_for_zip[lang][random_data_points[2]]

    doc_name = '{}.docx'.format(userInfo["userName"])
    doc.save(doc_name)
    print("Formatted And Saved Document with name {}".format(doc_name))
    subprocess.run(["libreoffice", "--headless", "--convert-to",
                   "pdf", "{}.docx".format(userInfo["userName"])])
    subprocess.run(
        ["lp", "-d", "myprinter", "{}.pdf".format(userInfo["userName"])])


if __name__ == "__main__":
    try:
        formatDocument({'userName': 'V', 'userId': 'd63142d7cf6f53a093ebc32ae1448f18', 'a1': '1', 'a2': '2',
                       'a3': '1', 'zipcode': '11222', 'sugarIntake': '3', 'archivePermission': '1', 'lang': 'en'})
    except Exception as e:
        print(e)
